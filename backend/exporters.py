"""导出工具：把调研大纲 / 最终报告生成为 Markdown 或 Word(.docx) 文件内容。

全部为纯函数，便于单元测试。Word 依赖 python-docx。

对外提供：
- outline_to_markdown(outline, topic)  大纲 → Markdown 文本
- outline_to_docx(outline, topic)      大纲 → docx 二进制
- report_to_markdown(text)             报告（本身是 Markdown）原样返回
- report_to_docx(text)                 报告 Markdown → docx 二进制
"""
import io
import re
from typing import Any, Dict, List

from pydantic import BaseModel


def _as_dict(obj: Any) -> Dict[str, Any]:
    """把 Pydantic 模型或 dict 统一成 dict（outline 可能来自 state 或 interrupt payload）。"""
    if isinstance(obj, BaseModel):
        return obj.model_dump()
    return dict(obj)


# ---------------------------------------------------------------------------
# 大纲导出
# ---------------------------------------------------------------------------
def outline_to_markdown(outline: Any, topic: str = "") -> str:
    """把大纲转成结构清晰的 Markdown 文本。"""
    o = _as_dict(outline)
    sections: List[str] = o.get("sections") or []
    questions: List[str] = o.get("key_questions") or []

    lines = ["# 调研大纲", ""]
    if topic:
        lines += [f"> 调研主题：{topic}", ""]
    lines += ["## 章节结构", ""]
    lines += [f"{i}. {s}" for i, s in enumerate(sections, 1)]
    lines += ["", "## 关键问题", ""]
    lines += [f"- {q}" for q in questions]
    return "\n".join(lines)


def outline_to_docx(outline: Any, topic: str = "") -> bytes:
    """把大纲转成 docx 二进制（标题 + 章节 + 关键问题）。"""
    from docx import Document

    o = _as_dict(outline)
    sections: List[str] = o.get("sections") or []
    questions: List[str] = o.get("key_questions") or []

    doc = Document()
    doc.add_heading("调研大纲", level=0)
    if topic:
        doc.add_paragraph(f"调研主题：{topic}")
    doc.add_heading("章节结构", level=1)
    for s in sections:
        doc.add_paragraph(str(s), style="List Number")
    doc.add_heading("关键问题", level=1)
    for q in questions:
        doc.add_paragraph(str(q), style="List Bullet")
    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# 报告导出
# ---------------------------------------------------------------------------
def report_to_markdown(text: str) -> str:
    """报告本身是 Markdown 文本，原样返回（保留原有格式）。"""
    return text


def report_to_docx(text: str) -> bytes:
    """把报告 Markdown 转成 docx（标题/列表/表格/段落）。"""
    from docx import Document

    doc = Document()
    _md_to_docx(doc, text)
    return _doc_to_bytes(doc)


def _doc_to_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Markdown → Word 解析
# ---------------------------------------------------------------------------
def _md_to_docx(doc, text: str) -> None:
    """把 Markdown 文本逐行写入 docx。

    支持：标题(#/##/###)、无序列表(-)、有序列表(1.)、表格(| a | b |)、水平线、普通段落。
    解析不了的表格降级为普通段落，保证不抛异常。
    """
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 表格：本行以 | 开头且下一行为分隔行
        if line.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            header = _split_row(line)
            i += 2
            rows: List[List[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i]))
                i += 1
            _add_table(doc, header, rows)
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # 无序列表
        m = re.match(r"^[-*]\s+(.*)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            i += 1
            continue

        # 有序列表
        m = re.match(r"^\d+[.)]\s+(.*)", line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Number")
            i += 1
            continue

        # 水平线（--- 等）
        if re.match(r"^\s*([-*_])\s*(\1\s*){2,}$", line):
            doc.add_paragraph()
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        doc.add_paragraph(line.strip())
        i += 1


def _split_row(line: str) -> List[str]:
    """把表格行拆成单元格（去掉首尾 |）。"""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_table_sep(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行（如 | --- | :---: |）。"""
    line = line.strip()
    if not line.startswith("|"):
        return False
    return all(re.fullmatch(r":?-{3,}:?", c) for c in _split_row(line))


def _add_table(doc, header: List[str], rows: List[List[str]]) -> None:
    """把表头与数据行写入 Word 表格；失败（列数不齐/样式缺失）时降级为段落。"""
    if not header:
        return
    ncol = max(len(header), max((len(r) for r in rows), default=0))
    if ncol == 0:
        return
    try:
        table = doc.add_table(rows=1, cols=ncol)
        table.style = "Table Grid"
        for j, cell in enumerate(header[:ncol]):
            table.rows[0].cells[j].text = cell
        for row in rows:
            cells = table.add_row().cells
            for j in range(ncol):
                cells[j].text = row[j] if j < len(row) else ""
    except Exception:  # noqa: BLE001 - 样式名缺失等，降级为普通文本行
        for row in [header] + rows:
            doc.add_paragraph(" | ".join(row))
