"""导出功能单元测试：exporters 模块（大纲/报告 → Markdown / Word）。

不依赖真实 LLM 或外部 API，可随时运行。
"""
import io

import pytest
from docx import Document

from exporters import (
    outline_to_docx,
    outline_to_markdown,
    report_to_docx,
    report_to_markdown,
)
from models import Outline


# ---------------------------------------------------------------------------
# 大纲 → Markdown
# ---------------------------------------------------------------------------
def test_outline_to_markdown_from_dict():
    md = outline_to_markdown(
        {"sections": ["行业概况", "竞争格局"], "key_questions": ["市场规模多少?"]},
        topic="咖啡机市场",
    )
    assert "# 调研大纲" in md
    assert "咖啡机市场" in md
    assert "1. 行业概况" in md and "2. 竞争格局" in md
    assert "- 市场规模多少?" in md


def test_outline_to_markdown_from_pydantic_model():
    o = Outline(sections=["行业概况"], key_questions=["市场规模?"])
    md = outline_to_markdown(o, topic="咖啡机市场")
    assert "行业概况" in md
    assert "市场规模?" in md


# ---------------------------------------------------------------------------
# 大纲 → Word
# ---------------------------------------------------------------------------
def test_outline_to_docx_valid_content():
    blob = outline_to_docx(
        {"sections": ["行业概况", "竞争格局"], "key_questions": ["市场规模?"]},
        topic="咖啡机市场",
    )
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs]
    assert "调研大纲" in texts
    assert "调研主题：咖啡机市场" in texts
    assert "行业概况" in texts and "竞争格局" in texts
    assert "市场规模?" in texts


# ---------------------------------------------------------------------------
# 报告 → Markdown（原样）
# ---------------------------------------------------------------------------
def test_report_to_markdown_identity():
    md = "# 报告\n\n- 要点"
    assert report_to_markdown(md) == md


# ---------------------------------------------------------------------------
# 报告 Markdown → Word
# ---------------------------------------------------------------------------
def test_report_to_docx_headings_and_lists():
    blob = report_to_docx("# 标题\n\n## 小节\n\n- 要点一\n- 要点二\n\n正文段落")
    doc = Document(io.BytesIO(blob))
    texts = [p.text for p in doc.paragraphs]
    assert "标题" in texts
    assert "小节" in texts
    assert "要点一" in texts and "要点二" in texts
    assert "正文段落" in texts


def test_report_to_docx_table():
    blob = report_to_docx("# 表格示例\n\n| 品牌 | 份额 |\n| --- | --- |\n| A | 30% |\n| B | 20% |")
    doc = Document(io.BytesIO(blob))
    assert doc.tables, "Markdown 表格应解析为 Word 表格"
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "品牌"
    assert table.rows[1].cells[1].text == "30%"


def test_report_to_docx_malformed_table_degrades():
    """列数不齐/缺表头等异常应降级为普通段落而不抛错。"""
    blob = report_to_docx("| a |\n| --- |\n| 1 |\n普通文本")
    doc = Document(io.BytesIO(blob))
    assert any("普通文本" in p.text for p in doc.paragraphs)
